"""Парсеры документов: из файла в набор блоков с привязкой к странице/секции."""

from __future__ import annotations

import re
from pathlib import Path
from typing import Protocol

from app.core.errors import ValidationError
from app.ingestion.models import ParsedBlock, ParsedDocument
from app.ingestion.normalization import normalize_text


class DocumentParser(Protocol):
    """Единый контракт парсера: подмена формата не влияет на остальной pipeline."""

    extensions: tuple[str, ...]

    def parse(self, path: Path) -> ParsedDocument: ...


class TextParser:
    extensions = (".txt",)

    def parse(self, path: Path) -> ParsedDocument:
        text = normalize_text(path.read_text(encoding="utf-8", errors="replace"))
        blocks = [
            ParsedBlock(text=paragraph)
            for paragraph in text.split("\n\n")
            if paragraph.strip()
        ]
        return ParsedDocument(blocks=blocks, meta={"format": "txt"})


class MarkdownParser:
    """Markdown: заголовки становятся секциями, чтобы citations были адресными."""

    extensions = (".md", ".markdown")
    _HEADING = re.compile(r"^(#{1,6})\s+(.*)$")

    def parse(self, path: Path) -> ParsedDocument:
        raw = path.read_text(encoding="utf-8", errors="replace")
        blocks: list[ParsedBlock] = []
        section: str | None = None
        buffer: list[str] = []

        def flush() -> None:
            if buffer:
                text = normalize_text("\n".join(buffer))
                if text:
                    blocks.append(ParsedBlock(text=text, section=section))
                buffer.clear()

        for line in raw.replace("\r\n", "\n").split("\n"):
            heading = self._HEADING.match(line)
            if heading:
                flush()
                section = heading.group(2).strip()
                continue
            if not line.strip():
                flush()
                continue
            buffer.append(line)
        flush()
        return ParsedDocument(blocks=blocks, meta={"format": "markdown"})


class PdfParser:
    """PDF: одна страница — один блок, номер страницы сохраняется для цитат."""

    extensions = (".pdf",)

    def parse(self, path: Path) -> ParsedDocument:
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        blocks: list[ParsedBlock] = []
        for number, page in enumerate(reader.pages, start=1):
            text = normalize_text(page.extract_text() or "")
            if text:
                blocks.append(ParsedBlock(text=text, page=number))
        return ParsedDocument(
            blocks=blocks, meta={"format": "pdf", "pages": str(len(reader.pages))}
        )


class DocxParser:
    """DOCX: заголовки Word (Heading 1-6) становятся секциями, как в Markdown."""

    extensions = (".docx",)
    _HEADING_STYLE = re.compile(r"^Heading (\d)$")

    def parse(self, path: Path) -> ParsedDocument:
        from docx import Document as WordDocument

        document = WordDocument(str(path))
        blocks: list[ParsedBlock] = []
        section: str | None = None

        for paragraph in document.paragraphs:
            raw = paragraph.text
            if not raw.strip():
                continue
            if self._HEADING_STYLE.match(paragraph.style.name if paragraph.style else ""):
                section = raw.strip()
                continue
            text = normalize_text(raw)
            if text:
                blocks.append(ParsedBlock(text=text, section=section))

        for table in document.tables:
            # Таблицы часто несут регламентные значения (сроки, ставки) —
            # тримминг ячеек через " | " сохраняет их читаемыми в чанке.
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    blocks.append(ParsedBlock(text=" | ".join(cells), section=section))

        return ParsedDocument(blocks=blocks, meta={"format": "docx"})


class HtmlParser:
    """HTML: h1-h6 становятся секциями, скрипты/стили/навигация отбрасываются."""

    extensions = (".html", ".htm")
    _HEADINGS = {"h1", "h2", "h3", "h4", "h5", "h6"}
    _SKIP_TAGS = {"script", "style", "nav", "header", "footer", "noscript"}
    _BLOCK_TAGS = {"p", "li", "td", "th", "blockquote", "pre", "div"}

    def parse(self, path: Path) -> ParsedDocument:
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(
            path.read_text(encoding="utf-8", errors="replace"), "lxml"
        )
        for tag in soup.find_all(self._SKIP_TAGS):
            tag.decompose()

        blocks: list[ParsedBlock] = []
        section: str | None = None
        seen: set[int] = set()

        for element in soup.find_all(self._HEADINGS | self._BLOCK_TAGS):
            if id(element) in seen or any(
                id(parent) in seen for parent in element.parents
            ):
                continue
            raw = element.get_text(" ", strip=True)
            if not raw:
                continue
            if element.name in self._HEADINGS:
                section = raw
                continue
            text = normalize_text(raw)
            if text:
                blocks.append(ParsedBlock(text=text, section=section))
                seen.add(id(element))

        return ParsedDocument(blocks=blocks, meta={"format": "html"})


class ParserRegistry:
    """Выбор парсера по расширению файла."""

    def __init__(self, parsers: list[DocumentParser] | None = None) -> None:
        self._by_extension: dict[str, DocumentParser] = {}
        for parser in parsers if parsers is not None else default_parsers():
            self.register(parser)

    def register(self, parser: DocumentParser) -> None:
        for extension in parser.extensions:
            self._by_extension[extension] = parser

    @property
    def supported_extensions(self) -> tuple[str, ...]:
        return tuple(sorted(self._by_extension))

    def parse(self, path: Path) -> ParsedDocument:
        parser = self._by_extension.get(path.suffix.lower())
        if parser is None:
            raise ValidationError(
                f"Формат {path.suffix or '<без расширения>'} не поддерживается. "
                f"Доступны: {', '.join(self.supported_extensions)}"
            )
        return parser.parse(path)


def default_parsers() -> list[DocumentParser]:
    return [
        TextParser(),
        MarkdownParser(),
        PdfParser(),
        DocxParser(),
        HtmlParser(),
    ]
