import pytest

from app.core.errors import ValidationError
from app.ingestion.parsers import ParserRegistry


@pytest.fixture
def registry() -> ParserRegistry:
    return ParserRegistry()


def test_parses_plain_text_into_paragraph_blocks(tmp_path, registry):
    path = tmp_path / "note.txt"
    path.write_text("первый абзац\n\nвторой абзац", encoding="utf-8")

    document = registry.parse(path)

    assert [block.text for block in document.blocks] == ["первый абзац", "второй абзац"]


def test_markdown_headings_become_sections(tmp_path, registry):
    path = tmp_path / "policy.md"
    path.write_text("# Отпуска\n\nПравила отпуска.\n\n## Перенос\n\nПравила переноса.",
                    encoding="utf-8")

    blocks = registry.parse(path).blocks

    assert [(b.section, b.text) for b in blocks] == [
        ("Отпуска", "Правила отпуска."),
        ("Перенос", "Правила переноса."),
    ]


def test_unsupported_format_is_rejected(tmp_path, registry):
    path = tmp_path / "archive.zip"
    path.write_bytes(b"PK")

    with pytest.raises(ValidationError, match="не поддерживается"):
        registry.parse(path)


def test_pdf_pages_are_preserved_for_citations(tmp_path, registry):
    pytest.importorskip("pypdf")
    from pypdf import PdfWriter

    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    path = tmp_path / "empty.pdf"
    with path.open("wb") as handle:
        writer.write(handle)

    document = registry.parse(path)

    assert document.meta["format"] == "pdf"
    assert document.meta["pages"] == "1"


def test_docx_headings_become_sections(tmp_path, registry):
    pytest.importorskip("docx")
    from docx import Document as WordDocument

    document = WordDocument()
    document.add_heading("Отпуска", level=1)
    document.add_paragraph("Отпуск предоставляется ежегодно.")
    document.add_heading("Перенос", level=2)
    document.add_paragraph("Перенос согласуется с руководителем.")
    path = tmp_path / "policy.docx"
    document.save(path)

    blocks = registry.parse(path).blocks

    assert [(b.section, b.text) for b in blocks] == [
        ("Отпуска", "Отпуск предоставляется ежегодно."),
        ("Перенос", "Перенос согласуется с руководителем."),
    ]


def test_docx_tables_are_flattened_into_readable_rows(tmp_path, registry):
    pytest.importorskip("docx")
    from docx import Document as WordDocument

    document = WordDocument()
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Ставка"
    table.rows[0].cells[1].text = "10%"
    path = tmp_path / "rates.docx"
    document.save(path)

    blocks = registry.parse(path).blocks

    assert blocks[0].text == "Ставка | 10%"


def test_html_headings_become_sections_and_scripts_are_dropped(tmp_path, registry):
    pytest.importorskip("bs4")
    html = """
    <html><body>
    <script>alert('x')</script>
    <h1>Отпуска</h1>
    <p>Отпуск предоставляется ежегодно.</p>
    <h2>Перенос</h2>
    <p>Перенос согласуется с руководителем.</p>
    </body></html>
    """
    path = tmp_path / "policy.html"
    path.write_text(html, encoding="utf-8")

    document = registry.parse(path)

    assert [(b.section, b.text) for b in document.blocks] == [
        ("Отпуска", "Отпуск предоставляется ежегодно."),
        ("Перенос", "Перенос согласуется с руководителем."),
    ]
    assert "alert" not in document.text
